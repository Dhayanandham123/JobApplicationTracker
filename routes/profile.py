import os
import io
import json
import pypdf
import docx
from flask import Blueprint, request, jsonify, session, g
from werkzeug.security import generate_password_hash, check_password_hash
from routes.auth import login_required
from database.db import (
    get_db, 
    get_user_by_id, 
    update_user_profile, 
    get_user_settings, 
    update_user_settings, 
    delete_user_account
)

profile_bp = Blueprint('profile', __name__)

@profile_bp.route('/api/profile', methods=['GET'])
@login_required
def get_profile():
    user_id = session['user_id']
    user = get_user_by_id(user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    user_dict = dict(user)
    user_dict.pop('password_hash', None)
    return jsonify(user_dict)

@profile_bp.route('/api/profile', methods=['PUT'])
@login_required
def update_profile():
    user_id = session['user_id']
    data = request.get_json() or request.form
    
    updated_user = update_user_profile(user_id, data)
    user_dict = dict(updated_user)
    user_dict.pop('password_hash', None)
    return jsonify({'success': True, 'user': user_dict})

def extract_text_from_file(file_storage):
    filename = file_storage.filename or ''
    ext = os.path.splitext(filename)[1].lower()
    
    file_bytes = file_storage.read()
    file_stream = io.BytesIO(file_bytes)
    
    extracted_text = ''
    
    if ext == '.pdf':
        try:
            reader = pypdf.PdfReader(file_stream)
            page_texts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    page_texts.append(text)
            extracted_text = '\n'.join(page_texts)
        except Exception as e:
            raise ValueError(f"Error parsing PDF file: {str(e)}")
            
    elif ext in ['.docx', '.doc']:
        try:
            doc = docx.Document(file_stream)
            full_text = []
            for para in doc.paragraphs:
                if para.text:
                    full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(' | '.join(row_text))
            extracted_text = '\n'.join(full_text)
        except Exception as e:
            raise ValueError(f"Error parsing Word document: {str(e)}")
            
    elif ext in ['.txt', '.rtf', '.md']:
        try:
            extracted_text = file_bytes.decode('utf-8', errors='ignore')
        except Exception as e:
            raise ValueError(f"Error reading text file: {str(e)}")
    else:
        raise ValueError("Unsupported file format. Please upload a PDF (.pdf), Word document (.docx), or text (.txt) file.")

    return filename, extracted_text.strip()

def recalculate_user_fit_scores(user_id, resume_text):
    db = get_db()
    apps = db.execute('SELECT * FROM applications WHERE user_id = ?', (user_id,)).fetchall()
    
    for app in apps:
        if not resume_text:
            db.execute('UPDATE applications SET fit_score = NULL, missing_skills = NULL WHERE id = ?', (app['id'],))
            continue
        jd_text = f"{app['company_name']} {app['job_title']}. {app['notes'] or ''}"
        try:
            fit_score, missing_skills = compute_fit_score(jd_text, resume_text)
            missing_skills_json = json.dumps(missing_skills) if missing_skills else None
            
            db.execute('''
                UPDATE applications
                SET fit_score = ?, missing_skills = ?
                WHERE id = ? AND user_id = ?
            ''', (fit_score, missing_skills_json, app['id'], user_id))
        except Exception as e:
            print(f"Error computing fit score for app {app['id']}: {e}")
    
    db.commit()
    return len(apps)

@profile_bp.route('/api/resume', methods=['GET'])
@login_required
def get_resume():
    user_id = session['user_id']
    user = get_user_by_id(user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
    return jsonify({
        'resume_text': user['resume_text'] or '',
        'resume_filename': user['resume_filename'] or ''
    })

@profile_bp.route('/api/resume/upload', methods=['POST'])
@login_required
def upload_resume():
    user_id = session['user_id']
    file_storage = request.files.get('resume_file') or request.files.get('file')
    
    if not file_storage or not file_storage.filename:
        return jsonify({'error': 'No file selected for upload.'}), 400
        
    try:
        filename, extracted_text = extract_text_from_file(file_storage)
    except ValueError as ve:
        return jsonify({'error': str(ve)}), 400
    except Exception as err:
        return jsonify({'error': f'Failed to process file: {str(err)}'}), 500
        
    if not extracted_text:
        return jsonify({'error': 'Could not extract text from this document. If it is a PDF, please ensure it contains selectable text (not a scanned image) or try a Word (.docx) file.'}), 400

    db = get_db()
    db.execute('UPDATE users SET resume_text = ?, resume_filename = ? WHERE id = ?', (extracted_text, filename, user_id))
    db.commit()

    # Recalculate fit scores for user applications
    apps_updated = recalculate_user_fit_scores(user_id, extracted_text)

    return jsonify({
        'success': True,
        'filename': filename,
        'resume_text': extracted_text,
        'apps_updated': apps_updated
    })

@profile_bp.route('/api/resume', methods=['DELETE'])
@login_required
def delete_resume():
    user_id = session['user_id']
    db = get_db()
    db.execute('UPDATE users SET resume_text = NULL, resume_filename = NULL WHERE id = ?', (user_id,))
    db.execute('UPDATE applications SET fit_score = NULL, missing_skills = NULL WHERE user_id = ?', (user_id,))
    db.commit()
    return jsonify({'success': True})

@profile_bp.route('/api/settings', methods=['GET'])
@login_required
def get_settings_route():
    user_id = session['user_id']
    settings = get_user_settings(user_id)
    return jsonify(settings)

@profile_bp.route('/api/settings', methods=['PUT'])
@login_required
def update_settings_route():
    user_id = session['user_id']
    data = request.get_json() or request.form
    
    updated = update_user_settings(user_id, data)
    return jsonify({'success': True, 'settings': updated})

@profile_bp.route('/api/account/change-password', methods=['POST'])
@login_required
def change_password():
    user_id = session['user_id']
    data = request.get_json() or request.form
    
    current_pw = data.get('current_password', '')
    new_pw = data.get('new_password', '')
    confirm_pw = data.get('confirm_password', '')
    
    if not current_pw or not new_pw:
        return jsonify({'error': 'Please provide current and new password.'}), 400
    
    if new_pw != confirm_pw:
        return jsonify({'error': 'New passwords do not match.'}), 400
        
    if len(new_pw) < 6:
        return jsonify({'error': 'New password must be at least 6 characters long.'}), 400
        
    user = get_user_by_id(user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
        
    if user['password_hash'] and not check_password_hash(user['password_hash'], current_pw):
        return jsonify({'error': 'Incorrect current password.'}), 400
        
    new_hash = generate_password_hash(new_pw)
    db = get_db()
    db.execute('UPDATE users SET password_hash = ? WHERE id = ?', (new_hash, user_id))
    db.commit()
    
    return jsonify({'success': True, 'message': 'Password updated successfully.'})

@profile_bp.route('/api/account/delete', methods=['POST'])
@login_required
def delete_account():
    user_id = session['user_id']
    delete_user_account(user_id)
    session.clear()
    return jsonify({'success': True, 'redirect': '/welcome'})
